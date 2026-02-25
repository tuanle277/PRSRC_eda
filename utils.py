"""
Shared utilities for PRSRC data work:
- Robust CSV reading
- DOCX yellow-highlight extraction
- Age-in-years QC helpers
- Null / mixed-type profiling
- Duplicate EncounterId checks
"""
import os
from pathlib import Path
from typing import List, Tuple

import pandas as pd
from docx import Document
from docx.enum.text import WD_COLOR_INDEX


def read_csv_robust(filepath: str) -> pd.DataFrame:
    """Robust CSV reader handling bad lines and mixed encodings."""
    try:
        return pd.read_csv(
            filepath,
            on_bad_lines="skip",
            engine="python",
            encoding="utf-8",
        )
    except UnicodeDecodeError:
        return pd.read_csv(
            filepath,
            on_bad_lines="skip",
            engine="python",
            encoding="latin-1",
        )
    except TypeError:
        # pandas < 2.0
        try:
            return pd.read_csv(
                filepath,
                error_bad_lines=False,  # type: ignore[arg-type]
                engine="python",
                encoding="utf-8",
            )
        except UnicodeDecodeError:
            return pd.read_csv(
                filepath,
                error_bad_lines=False,  # type: ignore[arg-type]
                engine="python",
                encoding="latin-1",
            )


# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------

def iter_paragraphs_and_cells(doc: Document):
    """
    Yield all paragraph objects from a python-docx Document, including those
    inside tables.
    """
    for p in doc.paragraphs:
        yield p

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def extract_yellow_highlights(docx_path: str) -> List[str]:
    """
    Return a list of text runs that are highlighted in yellow in a .docx file.
    """
    doc = Document(docx_path)
    hits: List[str] = []

    for para in iter_paragraphs_and_cells(doc):
        for run in para.runs:
            if run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                text = run.text
                if text and text.strip():
                    hits.append(text)

    return hits


def extract_yellow_highlights_to_file(docx_path: str) -> Path:
    """
    Extract yellow-highlighted runs from a .docx file and save them next to the
    document as '<original>_yellow_highlights.txt'. Returns the Path to the
    written text file.
    """
    path = Path(docx_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {docx_path}")

    hits = extract_yellow_highlights(str(path))

    out_path = path.with_suffix("")  # drop .docx
    out_path = out_path.with_name(out_path.name + "_yellow_highlights.txt")

    with out_path.open("w", encoding="utf-8") as f:
        for h in hits:
            f.write(h.replace("\n", " ") + "\n")

    print(f"Found {len(hits)} highlighted runs.")
    print(f"Saved to: {out_path}")

    return out_path


# ---------------------------------------------------------------------------
# Data QC helpers
# ---------------------------------------------------------------------------

def is_age_years_column(col_name: str) -> bool:
    """True if column name matches de-identified age-in-years columns (ex-BirthDate, ex-*Known)."""
    cn = col_name.strip().lower()
    if "date" not in cn or col_name == "BirthDate":
        return False
    if "known" in cn or cn.endswith("known"):
        return False
    return True


def count_negative_age_years_rows(
    age_years_path: str,
) -> List[Tuple[str, int, int, int]]:
    """
    Count how many rows have at least one negative value in de-identified age (years) columns
    across CSVs in a folder. Also reports total negative cell count.

    Returns list of (filename, rows_with_neg, negative_cells, total_rows).
    """
    path = os.path.abspath(age_years_path)
    if not os.path.isdir(path):
        print(f"Folder not found: {path}")
        return []

    csv_files = sorted([f for f in os.listdir(path) if f.endswith(".csv")])
    results: List[Tuple[str, int, int, int]] = []

    for filename in csv_files:
        filepath = os.path.join(path, filename)
        df = read_csv_robust(filepath)
        age_cols = [c for c in df.columns if is_age_years_column(c)]
        if not age_cols:
            continue

        has_neg = pd.DataFrame(
            {c: pd.to_numeric(df[c], errors="coerce") < 0 for c in age_cols}
        ).any(axis=1)
        rows_with_neg = int(has_neg.sum())
        neg_cells = int(
            sum((pd.to_numeric(df[c], errors="coerce") < 0).sum() for c in age_cols)
        )

        results.append((filename, rows_with_neg, neg_cells, len(df)))

    print("Negative de-identified age (years) rows in", path)
    print("-" * 70)
    if not results:
        print("No CSVs with age-year columns found.")
        return results
    for filename, n_rows, n_cells, total in results:
        row_label = "row" if n_rows == 1 else "rows"
        print(
            f"  {filename}: {n_rows} {row_label} with ≥1 negative age  |  "
            f"{n_cells} negative cells  |  {total} total rows"
        )
    print("-" * 70)
    print(f"  Total rows with ≥1 negative age: {sum(r[1] for r in results)}")
    print(f"  Total negative age cells: {sum(r[2] for r in results)}")
    return results


def profile_nulls_and_mixed_types(
    folder: str,
    output_csv: str | None = None,
    sample_non_null: int = 5000,
) -> pd.DataFrame:
    """
    For each CSV in a folder, compute per-column null counts and inferred Python
    value types, and flag columns with mixed types.
    """
    path = os.path.abspath(folder)
    if not os.path.isdir(path):
        print(f"Folder not found: {path}")
        return pd.DataFrame()

    csv_files = sorted([f for f in os.listdir(path) if f.endswith(".csv")])
    if not csv_files:
        print(f"No CSV files found in: {path}")
        return pd.DataFrame()

    rows: list[dict] = []

    for filename in csv_files:
        filepath = os.path.join(path, filename)
        df = read_csv_robust(filepath)
        n_rows = len(df)
        null_counts = df.isna().sum()

        for col in df.columns:
            s = df[col]
            null_count = int(null_counts.get(col, 0))
            dtype = str(s.dtype)

            non_null = s.dropna()
            if len(non_null) > sample_non_null:
                non_null = non_null.sample(sample_non_null, random_state=0)
            python_types = sorted({type(v).__name__ for v in non_null.tolist()}) if len(non_null) else []

            rows.append(
                {
                    "file": filename,
                    "column": col,
                    "rows": n_rows,
                    "dtype": dtype,
                    "null_count": null_count,
                    "null_pct": (null_count / n_rows) if n_rows else 0.0,
                    "python_types": "|".join(python_types),
                    "mixed_python_types": len(python_types) > 1,
                }
            )

    result = pd.DataFrame(rows)
    if output_csv is None:
        output_csv = os.path.join(path, "nulls_mixed_types_report.csv")
    if not result.empty:
        result.to_csv(output_csv, index=False)
        print(f"Wrote null/mixed-types report to {output_csv}")
    else:
        print("No data profiled; report is empty.")

    return result


def count_duplicate_encounter_ids(data_path: str) -> list[tuple[str, int, int, int]]:
    """
    Count duplicate EncounterId per CSV sheet in a folder.
    Returns list of (filename, n_duplicated_ids, n_duplicate_rows, total_rows).
    """
    data_dir = os.path.abspath(data_path)
    csv_files = sorted([f for f in os.listdir(data_dir) if f.endswith(".csv")])
    results: list[tuple[str, int, int, int]] = []

    for filename in csv_files:
        filepath = os.path.join(data_dir, filename)
        df = read_csv_robust(filepath)
        if "EncounterId" not in df.columns:
            continue
        vc = df["EncounterId"].value_counts()
        duplicated_ids = vc[vc > 1]
        n_duplicated_ids = int(len(duplicated_ids))
        n_duplicate_rows = int((vc[vc > 1] - 1).sum())  # extra rows per duplicated id
        results.append((filename, n_duplicated_ids, n_duplicate_rows, len(df)))

    print("EncounterId duplicate counts by sheet (in", data_path + ")")
    print("-" * 70)
    if not results:
        print("No CSVs with EncounterId column found.")
        return results
    for filename, n_dup_ids, n_dup_rows, total_rows in results:
        print(f"  {filename}")
        print(
            f"    EncounterIds that repeat: {n_dup_ids}  |  Extra rows (duplicates): "
            f"{n_dup_rows}  |  Total rows: {total_rows}"
        )
    print("-" * 70)
    total_dup_ids = sum(r[1] for r in results)
    total_dup_rows = sum(r[2] for r in results)
    print(f"  Total EncounterIds that repeat (across sheets): {total_dup_ids}")
    print(f"  Total duplicate rows (across sheets): {total_dup_rows}")
    return results

